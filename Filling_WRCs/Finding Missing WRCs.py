
import re
import string
import os
import errno

import pandas as pd
import numpy as np
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.datasets import fetch_20newsgroups
from sklearn.feature_extraction.text import CountVectorizer
from collections import Counter
from docx import Document
from docx.shared import Inches

stop_words  = set(stopwords.words('english'))

''' function to get only alphabets'''
def get_alpha(x):
    return re.sub(r'[^A-Za-z\s]+',' ',x)

''' function to remove punctuation'''
def remove_puntu(x):
    return x.translate(str.maketrans('', '', string.punctuation))

''' remove stopwords '''
def remove_stopwords(words):
    words =  [word for word in words if not word in stop_words]
    return words

''' remove stopwords '''
def remove_stopwords_line(words):
    words = words.split(" ")
    words =  [word for word in words if not word in stop_words]
    return " ".join(words)

''' find the difference between two list x and y '''
def find_exclusion(x,y):
    return list(set(y).difference(x))

''' find the intersection between two list x and y '''
def find_intersection(x,y):
    return list(set(x).intersection(y))

'''
In the tools given by onet there is a tag for each section of the tools which is removed by this function
'''
def remove_onet_tags(tools_list):
    final_list = dict()
    for val in tools_list:
        if len(val.split("-"))>1:
            key = val.split("-")[0]
            value = val.split("-")[1].strip()
        else:
            if val.find("-")!=-1:
                key = val.split("-")[0]
                value= ''
            else:
                key = 'N/A'
                value = val.split("-")[0].strip()
        final_list[key] = value.split(";")
    return final_list


'''
This functions fills the tools ,equipment section whenever its completely missing
Flow:
1. Find all the onet matched codes
2. Collate the tools and technology from the matchedones
3. Remove all the duplicates
4. Return the unique list
'''
def fill_missingtools(cta,onet,prediction):
    jsn = prediction['JSN']
    Onet_preds = prediction['Onet_preds'][:2]
    etmu = cta[cta.JSN==jsn]['EQUIPMENT, TOOLS, AND MATERIALS UTILIZED'].tolist()
    etmu = [x for x in etmu if pd.isna(x)!=True]
    etmu = [remove_puntu(x) for x in  etmu ]

    onet_tools = []
    onet_technology = []
    for val in Onet_preds:
        tools = onet[onet.Code==val]['Tools Used']
        technology = onet[onet.Code==val]['Technology Skills']
        if len(tools)>0:
            tools = tools.iloc[0]
            if pd.isna(tools)!=True:
                onet_tools.extend(tools.split("#"))
        if len(technology)>0:
            technology = technology.iloc[0]
            if pd.isna(technology)!=True:
                onet_technology.extend(technology.split("#"))
    tools_dict = remove_onet_tags(onet_tools)
    technology_dict = remove_onet_tags(onet_technology)
    tools_dict.update(technology_dict)

    return_tools = []
    for key in tools_dict.keys():
        sam_list = tools_dict.get(key)
        sam_list = list(set(sam_list).difference(etmu))
        return_tools.append(key+"-"+";".join(sam_list))
    return_tools = list(set(return_tools))
    return "#".join(return_tools)

'''
This function given any list and Number N will return N most commonly occuring items in the list
'''
def most_frequent(List,N):
    occurence_count = Counter(List)
    frequent = occurence_count.most_common(N)
    frequent = [x[0] for x in frequent]
    return frequent

'''
find the top 5 items in knowledge , skills and abilities with the matched onet codes
'''
def find_missingksa(cta,onet,prediction):
    jsn = prediction['JSN']
    Onet_preds = prediction['Onet_preds']
    k_list = []
    s_list = []
    a_list = []
    for val in Onet_preds:
        knowledge = onet[onet.Code==val]['Knowledge'].iloc[0]
        skills = onet[onet.Code==val]['Skills'].iloc[0]
        abilities = onet[onet.Code==val]['Abilities'].iloc[0]
        if pd.isna(knowledge)!=True:
            k_list.extend(knowledge.split("#"))
        if pd.isna(skills)!=True:
            s_list.extend(skills.split("#"))
        if pd.isna(abilities)!=True:
            a_list.extend(abilities.split("#"))
        k_list = most_frequent(k_list,5)
        s_list = most_frequent(a_list,5)
        a_list = most_frequent(a_list,5)

    return "#".join(k_list)+"#"+"#".join(s_list)+"#"+"#".join(a_list)

'''
Function will fill the education sections
'''
def fill_education(cta,onet,prediction):
    jsn = prediction['JSN']
    Onet_preds = prediction['Onet_preds']
    val_list = []
    for val in Onet_preds:
        education = onet[onet.Code==val]['Education'].iloc[0]
        if pd.isna(education)!=True:
            val_list.extend(education.split("#"))
    val_list = most_frequent(val_list,5)
    return "#".join(val_list)

def fill_primaryres(cta,onet,prediction):
    jsn = prediction['JSN']
    Onet_preds = prediction['Onet_preds']
    val_list = []
    for val in Onet_preds:
        Tasks = onet[onet.Code==val]['Tasks'].iloc[0]
        if pd.isna(Tasks)!=True:
            val_list.extend(Tasks.split("#"))
    val_list = most_frequent(val_list,10)
    return "#".join(val_list)

'''
This function will take all the updated job descriptions and creates single word document for each one of them
Flow:
1. Table is created at the top of the document with JSN , Job Title and
Update section:which will contain the section that is updated by this code
For example: if tools sections is filled by this code then update section will contain tools sections message
'''
def write_document(updated,updated_section,output_dir):
    document = Document()

    document.add_heading(updated['title'].iloc[0], 0)

    table = document.add_table(rows=3, cols=1)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'JSN: ' + str(updated['JSN'].iloc[0])
    hdr_cells = table.rows[1].cells
    hdr_cells[0].text = 'Job Title: ' + updated['title'].iloc[0]
    hdr_cells = table.rows[2].cells
    hdr_cells[0].text = 'Update Sections: ' + updated_section.title()

    column_names  = [ 'POSITION SUMMARY', 'PRIMARY RESPONSIBILITIES', 'EDUCATION/EXPERIENCE REQUIREMENTS', 'EQUIPMENT, TOOLS, AND MATERIALS UTILIZED', 'KNOWLEDGE, SKILLS, AND ABILITIES', 'MANAGEMENT RESPONSIBILITIES', 'PHYSICAL REQUIREMENTS', 'WORKING CONDITIONS','CHALLENGES']
    for val in column_names:
        document.add_heading(val, level=1)
        string_buf = updated[val].iloc[0]
        if pd.isna(string_buf)!= True:
            if val==updated_section:
                string_buf = string_buf.strip().split("#")
            else:
                string_buf = string_buf.strip().split(".")
            if len(string_buf)>1:
                string_buf = string_buf[:-1]
            for line in string_buf:
                document.add_paragraph(line, style='List Bullet')
        else:
            document.add_paragraph('')

    document.add_page_break()
    filename = updated['title'].iloc[0].replace("/","-")
    document.save(output_dir+filename+'.docx') # saves the word document

def main():
    parent_dir = os.path.dirname(os.getcwd())
    input_dir = os.path.join(parent_dir,"Modelling/input/")
    input_file = os.listdir(input_dir)
    if len(input_file)==0:
        raise FileNotFoundError(errno.ENOENT, os.strerror(errno.ENOENT), input_dir)
        sys.exit(0)
    cta = pd.read_csv(input_dir+input_file[0],sep="#")

    input_dir2 = os.path.join(parent_dir,"ONET SCRAPPING API/output/csv/")
    input_file2 = os.listdir(input_dir2)
    if len(input_file2)==0:
        raise FileNotFoundError(errno.ENOENT, os.strerror(errno.ENOENT), input_dir2)
        sys.exit(0)
    onet = pd.read_csv(input_dir2+input_file2[0])

    input_dir3 = os.path.join(parent_dir,"Modelling/output/")
    input_file3 = os.listdir(input_dir3)
    if len(input_file3)==0:
        raise FileNotFoundError(errno.ENOENT, os.strerror(errno.ENOENT), input_dir3)
        sys.exit(0)
    prediction = pd.read_csv(input_dir3+input_file3[0])

    prediction.columns = ['JSN', 'onet1', 'onet1_score', 'onet2', 'onet2_score', 'onet3',
       'onet3_score', 'onet4', 'onet4_score', 'onet5', 'onet5_score',
       'onet1_moc', 'onet2_moc', 'onet3_moc', 'onet4_moc', 'onet5_moc']

    prediction['JSN'] = prediction.JSN.apply(lambda x: int(x.split(" ")[1]))
    prediction['onet_code1'] = prediction.onet1.apply(lambda x: x.split(" ")[1])
    prediction['onet_code2'] = prediction.onet2.apply(lambda x: x.split(" ")[1])
    prediction['onet_code3'] = prediction.onet3.apply(lambda x: x.split(" ")[1])
    prediction['onet_code4'] = prediction.onet4.apply(lambda x: x.split(" ")[1])
    prediction['onet_code5'] = prediction.onet5.apply(lambda x: x.split(" ")[1])
    prediction['Onet_preds'] = prediction.apply(lambda x : [x['onet_code1'],x['onet_code2'],x['onet_code3'],x['onet_code4'],x['onet_code4'],x['onet_code5']],axis=1)
    prediction = prediction[['JSN','Onet_preds']]

    cta['PRIMARY RESPONSIBILITIES'] = cta['PRIMARY RESPONSIBILITIES'].apply(get_alpha)
    jsn_filled= []
    whatis_updated = []

    for index,row in cta.iterrows():

        if pd.isna(row['EQUIPMENT, TOOLS, AND MATERIALS UTILIZED'])==True:
            impute_val = fill_missingtools(cta,onet,{'JSN':row['JSN'],'Onet_preds':prediction[prediction.JSN==row['JSN']]['Onet_preds'].iloc[0]})
            if len(impute_val)>0:
                whatis_updated.append('EQUIPMENT, TOOLS, AND MATERIALS UTILIZED')
                jsn_filled.append(row['JSN'])
                cta.loc[index,'EQUIPMENT, TOOLS, AND MATERIALS UTILIZED'] = impute_val

        elif pd.isna(row['KNOWLEDGE, SKILLS, AND ABILITIES'])==True:
            impute_val = find_missingksa(cta,onet,{'JSN':row['JSN'],'Onet_preds':prediction[prediction.JSN==row['JSN']]['Onet_preds'].iloc[0]})
            if len(impute_val)>0:
                jsn_filled.append(row['JSN'])
                whatis_updated.append('KNOWLEDGE, SKILLS, AND ABILITIES')
            cta.loc[index,'KNOWLEDGE, SKILLS, AND ABILITIES'] = impute_val

        elif pd.isna(row['EDUCATION/EXPERIENCE REQUIREMENTS'])==True:
            impute_val = fill_education(cta,onet,{'JSN':row['JSN'],'Onet_preds':prediction[prediction.JSN==row['JSN']]['Onet_preds'].iloc[0]})
            if len(impute_val)>0:
                jsn_filled.append(row['JSN'])
                whatis_updated.append('EDUCATION/EXPERIENCE REQUIREMENTS')
            cta.loc[index,'EDUCATION/EXPERIENCE REQUIREMENTS'] = impute_val

        elif pd.isna(row['PRIMARY RESPONSIBILITIES'])==True:
            impute_val = fill_primaryres(cta,onet,{'JSN':row['JSN'],'Onet_preds':prediction[prediction.JSN==row['JSN']]['Onet_preds'].iloc[0]})
            if len(impute_val)>0:
                jsn_filled.append(row['JSN'])
                whatis_updated.append('PRIMARY RESPONSIBILITIES')
            cta.loc[index,'PRIMARY RESPONSIBILITIES'] = impute_val

    print("Number of Jobs filled to completion: ", len(jsn_filled))

    output_dir = os.path.join(os.getcwd(),"filled_doc/")
    for i,jsn in enumerate(jsn_filled):
        write_document(cta[cta.JSN==jsn],whatis_updated[i],output_dir)

if __name__=="__main__":main()
